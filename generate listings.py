import json
from docx import Document

# Load your JSON file
with open("dubai_today_bayut_listings.json", "r", encoding="utf-8") as f:
    listings = json.load(f)

# Categorize function
def categorize_listing(listing):
    for i, entry in enumerate(listing.get("top_3_prices", [])):
        if entry["link"] == listing["link"]:
            if entry["is_today_listing"] and i == 0:
                return "Top of Top 3 Today"
            elif entry["is_today_listing"]:
                return "In Top 3 Today"
            elif i == 0:
                return "Top of Top 3"
            else:
                return "In Top 3"
    return "Not in Top 3"

# Add categories
for listing in listings:
    listing["category"] = categorize_listing(listing)

# Sorting priority
priority_order = {
    "Top of Top 3 Today": 0,
    "In Top 3 Today": 1,
    "Top of Top 3": 2,
    "In Top 3": 3,
    "Not in Top 3": 4,
}

# Sort
sorted_listings = sorted(listings, key=lambda x: (priority_order.get(x["category"], 5), x["price_number"]))

# Explanation
category_explanations = {
    "Top of Top 3 Today": "🔥 Best priced property newly listed today in its building or area.",
    "In Top 3 Today": "✅ Among the top 3 lowest-priced new listings today.",
    "Top of Top 3": "💎 Best overall price (not newly listed) in top 3.",
    "In Top 3": "👍 Among the 3 best priced listings (not new).",
    "Not in Top 3": "⚠ Above the top 3 lowest-priced listings in the area."
}

# Group listings
grouped = {}
for listing in sorted_listings:
    grouped.setdefault(listing["category"], []).append(listing)

# Create Word document
doc = Document()
doc.add_heading("Sorted Property Listings (Human-Readable Format)", level=1)
doc.add_paragraph("Listings are grouped by how they rank in their area, especially if they're newly listed today.\n")

for category in priority_order:
    items = grouped.get(category, [])
    if items:
        doc.add_heading(category_explanations[category], level=2)
        for l in items:
            best = min(l["top_3_prices"], key=lambda x: x["price"])
            extra_line = f"\n💡 Best Deal in Area: {best['link']} (AED {best['price']})" if "Top 3" in category else ""
            doc.add_paragraph(
                f"{l['title']}\n"
                f"Price: {l['price']}\n"
                f"Location: {l['location']}\n"
                f"Bedrooms: {l['beds']}\n"
                f"Link: {l['link']}{extra_line}",
                style="List Bullet"
            )

# Save file
doc.save("Sorted_Listings_Human_Readable.docx")
print("✅ Done! Check the Word file: Sorted_Listings_Human_Readable.docx")