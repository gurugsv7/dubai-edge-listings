import json
from docx import Document

# Load your JSON file
with open("rC:\Users\gurug\OneDrive\Desktop\web apps\dubai-edge-listings\dubai_today_bayut_listings.json", "r", encoding="utf-8") as f:
    listings = json.load(f)

# Categorize each listing
def categorize_listing(listing):
    if listing.get("is_cheapest_in_building"):
        if listing.get("top_1_price", {}).get("is_today_listing", False):
            return "Top Deal Today"
        else:
            return "Top Deal (Not Today)"
    return "Other"

# Add categories to listings
for listing in listings:
    listing["category"] = categorize_listing(listing)

# Priority for sorting
priority_order = {
    "Top Deal Today": 0,
    "Top Deal (Not Today)": 1,
    "Other": 2,
}

# Sort listings by category priority and price
sorted_listings = sorted(
    listings,
    key=lambda x: (priority_order.get(x["category"], 3), x.get("price_number", float("inf")))
)

# ✅ Deduplicate by title or link
seen_titles = set()
seen_links = set()
unique_listings = []

for l in sorted_listings:
    title = l.get("title", "").strip().lower()
    link = l.get("link", "").strip().lower()
    if title in seen_titles or link in seen_links:
        continue  # Skip duplicates
    seen_titles.add(title)
    seen_links.add(link)
    unique_listings.append(l)

# Explanation for each category
category_explanations = {
    "Top Deal Today": "🔥 Best price in the building listed today",
    "Top Deal (Not Today)": "💎 Best price in the building (older listing)",
    "Other": "⚠ Not the cheapest in the building"
}

# Group listings by category
from collections import defaultdict
grouped = defaultdict(list)
for l in unique_listings:
    grouped[l["category"]].append(l)

# 📝 Create Word document
doc = Document()
doc.add_heading("Unique Best Listings (Human-Readable)", level=1)
doc.add_paragraph("Only the best and most unique listings per property are shown below.\n")

for category in priority_order:
    listings = grouped.get(category, [])
    if listings:
        doc.add_heading(category_explanations[category], level=2)
        for l in listings:
            top = l.get("top_1_price", {})
            extra_line = (
                f"\n💡 Lowest Price in Building: {top.get('link')} (AED {top.get('price')})"
                if top else ""
            )
            doc.add_paragraph(
                f"{l['title']}\n"
                f"Price: {l['price']}\n"
                f"Location: {l['location']}\n"
                f"Bedrooms: {l['beds']}\n"
                f"Link: {l['link']}{extra_line}",
                style="List Bullet"
            )

# Save document
doc.save("Unique_Top_Listings_By_Building.docx")
print("✅ Done! File saved as: Unique_Top_Listings_By_Building.docx")